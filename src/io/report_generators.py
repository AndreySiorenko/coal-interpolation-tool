"""
Report generation module for coal deposit interpolation results.

Provides comprehensive reporting capabilities:
- PDF reports with charts and tables
- HTML interactive reports
- Word/Excel template-based reports
"""

import pandas as pd
import numpy as np
from typing import Dict, Any, List, Tuple, Optional, Union
import warnings
import logging
from pathlib import Path
from abc import ABC, abstractmethod
from datetime import datetime
import base64
import io


class BaseReportGenerator(ABC):
    """Base class for report generators."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    @abstractmethod
    def generate_report(self, data: Dict[str, Any], 
                       output_path: str, 
                       template_path: Optional[str] = None,
                       **kwargs):
        """Generate report from data."""
        pass
    
    def _format_number(self, value: float, decimals: int = 3) -> str:
        """Format number with appropriate decimal places."""
        if pd.isna(value):
            return "N/A"
        return f"{value:.{decimals}f}"
    
    def _format_percentage(self, value: float, decimals: int = 1) -> str:
        """Format percentage value."""
        if pd.isna(value):
            return "N/A"
        return f"{value:.{decimals}f}%"


class PDFReportGenerator(BaseReportGenerator):
    """
    PDF report generator using matplotlib and reportlab.
    
    Creates professional PDF reports with charts, tables, and formatting.
    """
    
    def __init__(self):
        super().__init__()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available."""
        try:
            import matplotlib.pyplot as plt
            import matplotlib.backends.backend_pdf as pdf_backend
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            self.plt = plt
            self.pdf_backend = pdf_backend
            self.reportlab = {
                'SimpleDocTemplate': SimpleDocTemplate,
                'Paragraph': Paragraph,
                'Spacer': Spacer,
                'Table': Table,
                'TableStyle': TableStyle,
                'Image': Image,
                'getSampleStyleSheet': getSampleStyleSheet,
                'ParagraphStyle': ParagraphStyle,
                'letter': letter,
                'A4': A4,
                'inch': inch,
                'colors': colors
            }
            self.dependencies_available = True
            
        except ImportError:
            self.dependencies_available = False
            self.logger.warning("matplotlib or reportlab not available - PDF reports limited")
    
    def generate_report(self, data: Dict[str, Any],
                       output_path: str,
                       template_path: Optional[str] = None,
                       title: str = "Coal Deposit Interpolation Report",
                       **kwargs):
        """
        Generate PDF report.
        
        Args:
            data: Report data containing results, metadata, etc.
            output_path: Path for output PDF file
            template_path: Optional template file (not used in basic version)
            title: Report title
            **kwargs: Additional parameters
        """
        if not self.dependencies_available:
            raise ImportError("matplotlib and reportlab are required for PDF reports")
        
        self.logger.info(f"Generating PDF report: {output_path}")
        
        # Create PDF document
        doc = self.reportlab['SimpleDocTemplate'](
            output_path,
            pagesize=self.reportlab['A4'],
            topMargin=0.75 * self.reportlab['inch']
        )
        
        # Build report content
        story = []
        styles = self.reportlab['getSampleStyleSheet']()
        
        # Title
        title_style = self.reportlab['ParagraphStyle'](
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=self.reportlab['colors'].darkblue,
            alignment=1  # Center
        )
        story.append(self.reportlab['Paragraph'](title, title_style))
        story.append(self.reportlab['Spacer'](1, 0.3 * self.reportlab['inch']))
        
        # Report metadata
        story.append(self._create_metadata_section(data, styles))
        story.append(self.reportlab['Spacer'](1, 0.2 * self.reportlab['inch']))
        
        # Data summary
        if 'data_summary' in data:
            story.append(self._create_data_summary_section(data['data_summary'], styles))
            story.append(self.reportlab['Spacer'](1, 0.2 * self.reportlab['inch']))
        
        # Interpolation results
        if 'interpolation_results' in data:
            story.append(self._create_interpolation_results_section(data['interpolation_results'], styles))
            story.append(self.reportlab['Spacer'](1, 0.2 * self.reportlab['inch']))
        
        # Validation results
        if 'validation_results' in data:
            story.append(self._create_validation_results_section(data['validation_results'], styles))
            story.append(self.reportlab['Spacer'](1, 0.2 * self.reportlab['inch']))
        
        # Charts
        if 'charts' in data:
            story.extend(self._create_charts_section(data['charts'], styles))
        
        # Build PDF
        doc.build(story)
        self.logger.info("PDF report generated successfully")
    
    def _create_metadata_section(self, data: Dict[str, Any], styles) -> List:
        """Create metadata section."""
        elements = []
        
        # Section header
        elements.append(self.reportlab['Paragraph']("Report Information", styles['Heading2']))
        
        # Metadata table
        metadata = [
            ['Generated on:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Project:', data.get('project_name', 'Coal Deposit Analysis')],
            ['Method:', data.get('interpolation_method', 'N/A')],
            ['Data points:', str(data.get('data_count', 'N/A'))],
        ]
        
        if 'extent' in data:
            extent = data['extent']
            metadata.extend([
                ['X range:', f"{extent.get('x_min', 0):.2f} - {extent.get('x_max', 0):.2f}"],
                ['Y range:', f"{extent.get('y_min', 0):.2f} - {extent.get('y_max', 0):.2f}"],
            ])
        
        table = self.reportlab['Table'](metadata, colWidths=[2*self.reportlab['inch'], 3*self.reportlab['inch']])
        table.setStyle(self.reportlab['TableStyle']([
            ('BACKGROUND', (0, 0), (0, -1), self.reportlab['colors'].lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), self.reportlab['colors'].black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, self.reportlab['colors'].black)
        ]))
        
        elements.append(table)
        return elements
    
    def _create_data_summary_section(self, summary: Dict[str, Any], styles) -> List:
        """Create data summary section."""
        elements = []
        
        # Section header
        elements.append(self.reportlab['Paragraph']("Data Summary", styles['Heading2']))
        
        # Summary statistics table
        stats_data = [['Statistic', 'Value']]
        
        if 'statistics' in summary:
            stats = summary['statistics']
            stats_data.extend([
                ['Mean', self._format_number(stats.get('mean'))],
                ['Median', self._format_number(stats.get('median'))],
                ['Std Dev', self._format_number(stats.get('std'))],
                ['Min', self._format_number(stats.get('min'))],
                ['Max', self._format_number(stats.get('max'))],
                ['Count', str(stats.get('count', 'N/A'))]
            ])
        
        table = self.reportlab['Table'](stats_data, colWidths=[2*self.reportlab['inch'], 2*self.reportlab['inch']])
        table.setStyle(self.reportlab['TableStyle']([
            ('BACKGROUND', (0, 0), (-1, 0), self.reportlab['colors'].grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), self.reportlab['colors'].whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, self.reportlab['colors'].black)
        ]))
        
        elements.append(table)
        return elements
    
    def _create_interpolation_results_section(self, results: Dict[str, Any], styles) -> List:
        """Create interpolation results section."""
        elements = []
        
        # Section header
        elements.append(self.reportlab['Paragraph']("Interpolation Results", styles['Heading2']))
        
        # Method details
        if 'method_info' in results:
            method_info = results['method_info']
            elements.append(self.reportlab['Paragraph'](
                f"<b>Method:</b> {method_info.get('name', 'Unknown')}", 
                styles['Normal']
            ))
            
            if 'parameters' in method_info:
                params_text = "Parameters: " + ", ".join([
                    f"{k}={v}" for k, v in method_info['parameters'].items()
                ])
                elements.append(self.reportlab['Paragraph'](params_text, styles['Normal']))
        
        # Quality metrics
        if 'quality_metrics' in results:
            elements.append(self.reportlab['Spacer'](1, 0.1 * self.reportlab['inch']))
            elements.append(self.reportlab['Paragraph']("Quality Metrics", styles['Heading3']))
            
            metrics = results['quality_metrics']
            metrics_data = [['Metric', 'Value']]
            
            for metric, value in metrics.items():
                if isinstance(value, (int, float)):
                    formatted_value = self._format_number(value)
                else:
                    formatted_value = str(value)
                metrics_data.append([metric.replace('_', ' ').title(), formatted_value])
            
            table = self.reportlab['Table'](metrics_data, colWidths=[2.5*self.reportlab['inch'], 1.5*self.reportlab['inch']])
            table.setStyle(self.reportlab['TableStyle']([
                ('BACKGROUND', (0, 0), (-1, 0), self.reportlab['colors'].darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), self.reportlab['colors'].whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, self.reportlab['colors'].black)
            ]))
            
            elements.append(table)
        
        return elements
    
    def _create_validation_results_section(self, validation: Dict[str, Any], styles) -> List:
        """Create validation results section."""
        elements = []
        
        # Section header
        elements.append(self.reportlab['Paragraph']("Validation Results", styles['Heading2']))
        
        if 'cross_validation' in validation:
            cv_results = validation['cross_validation']
            elements.append(self.reportlab['Paragraph']("Cross-Validation", styles['Heading3']))
            
            cv_data = [['Method', 'RMSE', 'R²', 'MAE']]
            for method_name, metrics in cv_results.items():
                cv_data.append([
                    method_name,
                    self._format_number(metrics.get('rmse', 0)),
                    self._format_number(metrics.get('r_squared', 0)),
                    self._format_number(metrics.get('mae', 0))
                ])
            
            table = self.reportlab['Table'](cv_data)
            table.setStyle(self.reportlab['TableStyle']([
                ('BACKGROUND', (0, 0), (-1, 0), self.reportlab['colors'].lightblue),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, self.reportlab['colors'].black)
            ]))
            
            elements.append(table)
        
        return elements
    
    def _create_charts_section(self, charts: List[str], styles) -> List:
        """Create charts section from image paths."""
        elements = []
        
        elements.append(self.reportlab['Paragraph']("Charts and Visualizations", styles['Heading2']))
        
        for chart_path in charts:
            if Path(chart_path).exists():
                # Add chart image
                img = self.reportlab['Image'](chart_path, width=6*self.reportlab['inch'], height=4*self.reportlab['inch'])
                elements.append(img)
                elements.append(self.reportlab['Spacer'](1, 0.2 * self.reportlab['inch']))
        
        return elements


class HTMLReportGenerator(BaseReportGenerator):
    """
    HTML report generator with interactive elements.
    
    Creates responsive HTML reports with charts, tables, and interactivity.
    """
    
    def __init__(self):
        super().__init__()
    
    def generate_report(self, data: Dict[str, Any],
                       output_path: str,
                       template_path: Optional[str] = None,
                       title: str = "Coal Deposit Interpolation Report",
                       **kwargs):
        """
        Generate HTML report.
        
        Args:
            data: Report data
            output_path: Path for output HTML file
            template_path: Optional template file
            title: Report title
            **kwargs: Additional parameters
        """
        self.logger.info(f"Generating HTML report: {output_path}")
        
        # Build HTML content
        html_content = self._build_html_report(data, title)
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.logger.info("HTML report generated successfully")
    
    def _build_html_report(self, data: Dict[str, Any], title: str) -> str:
        """Build complete HTML report."""
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        {self._get_css_styles()}
    </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/3.9.1/chart.min.js"></script>
</head>
<body>
    <div class="container">
        <header>
            <h1>{title}</h1>
            <p class="subtitle">Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </header>
        
        <main>
            {self._build_metadata_section(data)}
            {self._build_data_summary_section(data.get('data_summary', {}))}
            {self._build_interpolation_section(data.get('interpolation_results', {}))}
            {self._build_validation_section(data.get('validation_results', {}))}
            {self._build_charts_section(data.get('charts', []))}
        </main>
        
        <footer>
            <p>Report generated by Coal Deposit Interpolation Tool</p>
        </footer>
    </div>
    
    <script>
        {self._get_javascript()}
    </script>
</body>
</html>"""
        
        return html
    
    def _get_css_styles(self) -> str:
        """Get CSS styles for HTML report."""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f5f5f5;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        
        header {
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #2c5aa0;
        }
        
        h1 {
            color: #2c5aa0;
            font-size: 2.5em;
            margin-bottom: 10px;
        }
        
        .subtitle {
            color: #666;
            font-style: italic;
        }
        
        h2 {
            color: #2c5aa0;
            margin: 30px 0 15px 0;
            padding-bottom: 5px;
            border-bottom: 1px solid #ddd;
        }
        
        h3 {
            color: #4a4a4a;
            margin: 20px 0 10px 0;
        }
        
        .section {
            margin-bottom: 30px;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        
        th {
            background-color: #2c5aa0;
            color: white;
            font-weight: bold;
        }
        
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        
        .metrics-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }
        
        .metric-card {
            background: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            border-left: 4px solid #2c5aa0;
        }
        
        .metric-value {
            font-size: 1.5em;
            font-weight: bold;
            color: #2c5aa0;
        }
        
        .metric-label {
            color: #666;
            margin-top: 5px;
        }
        
        .chart-container {
            margin: 20px 0;
            padding: 20px;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        footer {
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
        }
        
        @media (max-width: 768px) {
            .container {
                padding: 10px;
            }
            
            h1 {
                font-size: 2em;
            }
            
            .metrics-grid {
                grid-template-columns: 1fr;
            }
        }
        """
    
    def _build_metadata_section(self, data: Dict[str, Any]) -> str:
        """Build metadata section HTML."""
        html = """
        <section class="section">
            <h2>Report Information</h2>
            <table>
                <tr><th>Property</th><th>Value</th></tr>
        """
        
        metadata = [
            ('Project', data.get('project_name', 'Coal Deposit Analysis')),
            ('Method', data.get('interpolation_method', 'N/A')),
            ('Data Points', str(data.get('data_count', 'N/A'))),
        ]
        
        if 'extent' in data:
            extent = data['extent']
            metadata.extend([
                ('X Range', f"{extent.get('x_min', 0):.2f} - {extent.get('x_max', 0):.2f}"),
                ('Y Range', f"{extent.get('y_min', 0):.2f} - {extent.get('y_max', 0):.2f}"),
            ])
        
        for key, value in metadata:
            html += f"<tr><td>{key}</td><td>{value}</td></tr>"
        
        html += """
            </table>
        </section>
        """
        
        return html
    
    def _build_data_summary_section(self, summary: Dict[str, Any]) -> str:
        """Build data summary section HTML."""
        if not summary:
            return ""
        
        html = """
        <section class="section">
            <h2>Data Summary</h2>
        """
        
        if 'statistics' in summary:
            stats = summary['statistics']
            html += '<div class="metrics-grid">'
            
            metrics = [
                ('Mean', self._format_number(stats.get('mean'))),
                ('Median', self._format_number(stats.get('median'))),
                ('Std Dev', self._format_number(stats.get('std'))),
                ('Min', self._format_number(stats.get('min'))),
                ('Max', self._format_number(stats.get('max'))),
                ('Count', str(stats.get('count', 'N/A')))
            ]
            
            for label, value in metrics:
                html += f"""
                <div class="metric-card">
                    <div class="metric-value">{value}</div>
                    <div class="metric-label">{label}</div>
                </div>
                """
            
            html += '</div>'
        
        html += '</section>'
        return html
    
    def _build_interpolation_section(self, results: Dict[str, Any]) -> str:
        """Build interpolation results section HTML."""
        if not results:
            return ""
        
        html = """
        <section class="section">
            <h2>Interpolation Results</h2>
        """
        
        # Method info
        if 'method_info' in results:
            method_info = results['method_info']
            html += f"<p><strong>Method:</strong> {method_info.get('name', 'Unknown')}</p>"
            
            if 'parameters' in method_info:
                html += "<p><strong>Parameters:</strong> "
                params = [f"{k}={v}" for k, v in method_info['parameters'].items()]
                html += ", ".join(params) + "</p>"
        
        # Quality metrics
        if 'quality_metrics' in results:
            html += "<h3>Quality Metrics</h3>"
            html += "<table><tr><th>Metric</th><th>Value</th></tr>"
            
            for metric, value in results['quality_metrics'].items():
                if isinstance(value, (int, float)):
                    formatted_value = self._format_number(value)
                else:
                    formatted_value = str(value)
                html += f"<tr><td>{metric.replace('_', ' ').title()}</td><td>{formatted_value}</td></tr>"
            
            html += "</table>"
        
        html += "</section>"
        return html
    
    def _build_validation_section(self, validation: Dict[str, Any]) -> str:
        """Build validation results section HTML."""
        if not validation:
            return ""
        
        html = """
        <section class="section">
            <h2>Validation Results</h2>
        """
        
        if 'cross_validation' in validation:
            html += "<h3>Cross-Validation</h3>"
            html += "<table><tr><th>Method</th><th>RMSE</th><th>R²</th><th>MAE</th></tr>"
            
            for method_name, metrics in validation['cross_validation'].items():
                html += f"""
                <tr>
                    <td>{method_name}</td>
                    <td>{self._format_number(metrics.get('rmse', 0))}</td>
                    <td>{self._format_number(metrics.get('r_squared', 0))}</td>
                    <td>{self._format_number(metrics.get('mae', 0))}</td>
                </tr>
                """
            
            html += "</table>"
        
        html += "</section>"
        return html
    
    def _build_charts_section(self, charts: List[str]) -> str:
        """Build charts section HTML."""
        if not charts:
            return ""
        
        html = """
        <section class="section">
            <h2>Charts and Visualizations</h2>
        """
        
        for i, chart_path in enumerate(charts):
            if Path(chart_path).exists():
                # Convert image to base64 for embedding
                try:
                    with open(chart_path, 'rb') as f:
                        image_data = base64.b64encode(f.read()).decode()
                    
                    ext = Path(chart_path).suffix.lower()
                    mime_type = 'image/png' if ext == '.png' else 'image/jpeg'
                    
                    html += f"""
                    <div class="chart-container">
                        <img src="data:{mime_type};base64,{image_data}" 
                             alt="Chart {i+1}" style="max-width: 100%; height: auto;">
                    </div>
                    """
                except Exception as e:
                    self.logger.warning(f"Could not embed chart {chart_path}: {e}")
        
        html += "</section>"
        return html
    
    def _get_javascript(self) -> str:
        """Get JavaScript for interactive elements."""
        return """
        // Add any interactive JavaScript here
        console.log('Coal Deposit Interpolation Report loaded');
        
        // Add click handlers for interactive elements
        document.addEventListener('DOMContentLoaded', function() {
            // Example: Add table sorting, chart interactions, etc.
        });
        """


class OfficeReportGenerator(BaseReportGenerator):
    """
    Word/Excel report generator using templates.
    
    Creates reports using Word/Excel templates with data substitution.
    """
    
    def __init__(self):
        super().__init__()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available."""
        try:
            import openpyxl
            from docx import Document
            
            self.openpyxl = openpyxl
            self.Document = Document
            self.dependencies_available = True
            
        except ImportError:
            self.dependencies_available = False
            self.logger.warning("openpyxl or python-docx not available - Office reports limited")
    
    def generate_report(self, data: Dict[str, Any],
                       output_path: str,
                       template_path: Optional[str] = None,
                       report_type: str = 'excel',
                       **kwargs):
        """
        Generate Office report.
        
        Args:
            data: Report data
            output_path: Path for output file
            template_path: Optional template file
            report_type: 'excel' or 'word'
            **kwargs: Additional parameters
        """
        if not self.dependencies_available:
            raise ImportError("openpyxl and python-docx are required for Office reports")
        
        self.logger.info(f"Generating {report_type} report: {output_path}")
        
        if report_type.lower() == 'excel':
            self._generate_excel_report(data, output_path, template_path)
        elif report_type.lower() == 'word':
            self._generate_word_report(data, output_path, template_path)
        else:
            raise ValueError(f"Unsupported report type: {report_type}")
        
        self.logger.info(f"{report_type.title()} report generated successfully")
    
    def _generate_excel_report(self, data: Dict[str, Any], 
                              output_path: str, 
                              template_path: Optional[str]):
        """Generate Excel report."""
        
        if template_path and Path(template_path).exists():
            # Use template
            wb = self.openpyxl.load_workbook(template_path)
        else:
            # Create new workbook
            wb = self.openpyxl.Workbook()
            
            # Remove default sheet
            if 'Sheet' in wb.sheetnames:
                wb.remove(wb['Sheet'])
        
        # Summary sheet
        if 'Summary' not in wb.sheetnames:
            ws_summary = wb.create_sheet('Summary')
        else:
            ws_summary = wb['Summary']
        
        self._populate_excel_summary(ws_summary, data)
        
        # Data sheet
        if 'data_summary' in data and 'raw_data' in data['data_summary']:
            if 'Data' not in wb.sheetnames:
                ws_data = wb.create_sheet('Data')
            else:
                ws_data = wb['Data']
            
            self._populate_excel_data(ws_data, data['data_summary']['raw_data'])
        
        # Validation sheet
        if 'validation_results' in data:
            if 'Validation' not in wb.sheetnames:
                ws_validation = wb.create_sheet('Validation')
            else:
                ws_validation = wb['Validation']
            
            self._populate_excel_validation(ws_validation, data['validation_results'])
        
        # Save workbook
        wb.save(output_path)
    
    def _populate_excel_summary(self, ws, data: Dict[str, Any]):
        """Populate Excel summary sheet."""
        ws['A1'] = 'Coal Deposit Interpolation Report'
        ws['A1'].font = self.openpyxl.styles.Font(size=16, bold=True)
        
        row = 3
        
        # Report info
        ws[f'A{row}'] = 'Report Information'
        ws[f'A{row}'].font = self.openpyxl.styles.Font(bold=True)
        row += 1
        
        info_data = [
            ('Generated on:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Project:', data.get('project_name', 'Coal Deposit Analysis')),
            ('Method:', data.get('interpolation_method', 'N/A')),
            ('Data points:', str(data.get('data_count', 'N/A'))),
        ]
        
        for label, value in info_data:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = value
            row += 1
        
        row += 2
        
        # Statistics
        if 'data_summary' in data and 'statistics' in data['data_summary']:
            ws[f'A{row}'] = 'Data Statistics'
            ws[f'A{row}'].font = self.openpyxl.styles.Font(bold=True)
            row += 1
            
            stats = data['data_summary']['statistics']
            for stat_name, value in stats.items():
                ws[f'A{row}'] = stat_name.replace('_', ' ').title()
                ws[f'B{row}'] = value if not pd.isna(value) else 'N/A'
                row += 1
    
    def _populate_excel_data(self, ws, data: pd.DataFrame):
        """Populate Excel data sheet."""
        # Headers
        for col_idx, column in enumerate(data.columns, 1):
            ws.cell(row=1, column=col_idx, value=column)
            ws.cell(row=1, column=col_idx).font = self.openpyxl.styles.Font(bold=True)
        
        # Data
        for row_idx, (_, row) in enumerate(data.iterrows(), 2):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
    
    def _populate_excel_validation(self, ws, validation: Dict[str, Any]):
        """Populate Excel validation sheet."""
        ws['A1'] = 'Validation Results'
        ws['A1'].font = self.openpyxl.styles.Font(size=14, bold=True)
        
        row = 3
        
        if 'cross_validation' in validation:
            ws[f'A{row}'] = 'Cross-Validation Results'
            ws[f'A{row}'].font = self.openpyxl.styles.Font(bold=True)
            row += 1
            
            # Headers
            headers = ['Method', 'RMSE', 'R²', 'MAE']
            for col_idx, header in enumerate(headers, 1):
                ws.cell(row=row, column=col_idx, value=header)
                ws.cell(row=row, column=col_idx).font = self.openpyxl.styles.Font(bold=True)
            row += 1
            
            # Data
            for method_name, metrics in validation['cross_validation'].items():
                ws.cell(row=row, column=1, value=method_name)
                ws.cell(row=row, column=2, value=metrics.get('rmse', 'N/A'))
                ws.cell(row=row, column=3, value=metrics.get('r_squared', 'N/A'))
                ws.cell(row=row, column=4, value=metrics.get('mae', 'N/A'))
                row += 1
    
    def _generate_word_report(self, data: Dict[str, Any],
                             output_path: str,
                             template_path: Optional[str]):
        """Generate Word report."""
        
        if template_path and Path(template_path).exists():
            # Use template
            doc = self.Document(template_path)
        else:
            # Create new document
            doc = self.Document()
        
        # Title
        title = doc.add_heading('Coal Deposit Interpolation Report', 0)
        
        # Report info
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = 'Property'
        hdr_cells[1].text = 'Value'
        
        info_data = [
            ('Generated on', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Project', data.get('project_name', 'Coal Deposit Analysis')),
            ('Method', data.get('interpolation_method', 'N/A')),
            ('Data points', str(data.get('data_count', 'N/A'))),
        ]
        
        for label, value in info_data:
            row_cells = info_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        # Data summary
        if 'data_summary' in data and 'statistics' in data['data_summary']:
            doc.add_heading('Data Summary', level=1)
            
            stats = data['data_summary']['statistics']
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Statistic'
            hdr_cells[1].text = 'Value'
            
            for stat_name, value in stats.items():
                row_cells = stats_table.add_row().cells
                row_cells[0].text = stat_name.replace('_', ' ').title()
                row_cells[1].text = self._format_number(value) if not pd.isna(value) else 'N/A'
        
        # Save document
        doc.save(output_path)


# Factory function
def create_report_generator(format_type: str, **kwargs) -> BaseReportGenerator:
    """
    Create appropriate report generator based on format.
    
    Args:
        format_type: Report format ('pdf', 'html', 'excel', 'word')
        **kwargs: Additional parameters
        
    Returns:
        Appropriate report generator instance
    """
    format_type = format_type.lower()
    
    if format_type == 'pdf':
        return PDFReportGenerator(**kwargs)
    elif format_type == 'html':
        return HTMLReportGenerator(**kwargs)
    elif format_type in ['excel', 'word']:
        return OfficeReportGenerator(**kwargs)
    else:
        raise ValueError(f"Unsupported report format: {format_type}")


def generate_report(data: Dict[str, Any],
                   output_path: str,
                   format_type: str,
                   **kwargs):
    """
    Convenience function to generate report in any supported format.
    
    Args:
        data: Report data
        output_path: Output file path
        format_type: Report format type
        **kwargs: Additional parameters
    """
    generator = create_report_generator(format_type, **kwargs)
    generator.generate_report(data, output_path, **kwargs)